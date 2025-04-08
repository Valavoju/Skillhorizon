from flask import Flask, jsonify, request
from flask_cors import CORS
from flask_jwt_extended import JWTManager, create_access_token, jwt_required
from pymongo import MongoClient
from werkzeug.security import generate_password_hash, check_password_hash
import PyPDF2
import docx
from dotenv import dotenv_values
from bson import ObjectId

# 🔹 Load environment variables from .env
config = dotenv_values(".env")
MONGO_URI = config["MONGO_URI"]
JWT_SECRET = config["JWT_SECRET"]

# 🔹 App Initialization
app = Flask(__name__)
CORS(app, resources={r"/*": {"origins": "http://localhost:3000"}})
app.config["JWT_SECRET_KEY"] = JWT_SECRET

# 🔹 MongoDB Setup
client = MongoClient(MONGO_URI)
db = client["skillhorizon"]
users_collection = db["users"]
upload_collection = db["resumes"]
job_collection = db["jobs"]

# ✅ Confirm connection
try:
    client.server_info()
    print("✅ Connected to MongoDB via dotenv_values")
except Exception as e:
    print("❌ MongoDB connection error:", e)

# 🔹 JWT Setup
jwt = JWTManager(app)

# 🔹 Utility Functions
def compute_match_score(resume, job):
    resume_skills = set(map(str.lower, resume["skills"]))
    job_skills = set(map(str.lower, job["Required Skills"].split(",")))

    skill_match_score = len(resume_skills & job_skills) / len(job_skills) * 100
    experience_match = 100 if resume["experience"] >= int(job["Experience Level"].split()[0]) else 50
    project_match = any(proj.lower() in job["Job Title"].lower() for proj in resume["projects"])

    final_score = (0.6 * skill_match_score) + (0.3 * experience_match) + (0.1 * (100 if project_match else 0))
    return round(final_score, 2)

def gap_analysis(resume, job):
    resume_skills = set(map(str.lower, resume["skills"]))
    job_skills = set(map(str.lower, job["Required Skills"].split(",")))

    return list(job_skills - resume_skills)

def match_resume_with_jobs(resume_id):
    resume = upload_collection.find_one({"_id": ObjectId(resume_id)})
    jobs = list(job_collection.find({}))

    results = []
    for job in jobs:
        score = compute_match_score(resume, job)
        missing_skills = gap_analysis(resume, job)
        results.append({
            "Job Title": job["Job Title"],
            "Match Score": score,
            "Missing Skills": missing_skills
        })
    return results

# 🔹 Routes
@app.route("/")
def home():
    return jsonify({"message": "Welcome to Skill Horizon Backend!"})

@app.route("/signup", methods=["POST"])
def sign_up():
    try:
        data = request.get_json()
        name = data.get("name")
        email = data.get("email")
        password = data.get("password")

        if not name or not email or not password:
            return jsonify({"error": "All fields required"}), 400

        if users_collection.find_one({"email": email}):
            return jsonify({"error": "User already exists"}), 400

        hashed_password = generate_password_hash(password)
        users_collection.insert_one({
            "name": name, "email": email, "password": hashed_password
        })

        return jsonify({"message": "User registered successfully!"}), 201

    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route("/login", methods=["POST"])
def login():
    try:
        data = request.get_json()
        email = data.get("email")
        password = data.get("password")

        user = users_collection.find_one({"email": email})
        if not user or not check_password_hash(user["password"], password):
            return jsonify({"error": "Invalid email or password"}), 401

        token = create_access_token(identity=email)
        return jsonify({"message": "Login successful", "access_token": token}), 200

    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route("/protected", methods=["GET"])
@jwt_required()
def protected():
    return jsonify({"message": "This is a protected route!"})

@app.route("/upload", methods=["POST"])
def upload():
    try:
        if "upload" not in request.files:
            return jsonify({"error": "No file uploaded"}), 400

        file = request.files["upload"]
        file_ext = file.filename.split(".")[-1].lower()

        if file_ext not in ["pdf", "docx"]:
            return jsonify({"error": "Only PDF and DOCX allowed"}), 400

        extracted_text = ""
        if file_ext == "pdf":
            reader = PyPDF2.PdfReader(file)
            for page in reader.pages:
                extracted_text += page.extract_text() + "\n"
        elif file_ext == "docx":
            doc = docx.Document(file)
            extracted_text = "\n".join([para.text for para in doc.paragraphs])

        result = upload_collection.insert_one({
            "file_name": file.filename,
            "content": extracted_text
        })

        return jsonify({"message": "File uploaded and stored successfully!", "resume_id": str(result.inserted_id)}), 201

    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route("/extract-text", methods=["POST"])
@jwt_required()
def extract_text():
    try:
        if "resume" not in request.files:
            return jsonify({"error": "No file uploaded"}), 400

        file = request.files["resume"]
        file_ext = file.filename.split(".")[-1].lower()

        if file_ext not in ["pdf", "docx"]:
            return jsonify({"error": "Only PDF and DOCX supported"}), 400

        extracted_text = ""
        if file_ext == "pdf":
            reader = PyPDF2.PdfReader(file)
            for page in reader.pages:
                extracted_text += page.extract_text() + "\n"
        elif file_ext == "docx":
            doc = docx.Document(file)
            extracted_text = "\n".join([para.text for para in doc.paragraphs])

        return jsonify({"message": "Text extracted successfully", "extractedText": extracted_text}), 200

    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route("/match-jobs", methods=["POST"])
@jwt_required()
def match_jobs():
    try:
        data = request.get_json()
        resume_id = data.get("resume_id")

        if not resume_id:
            return jsonify({"error": "Resume ID is required"}), 400

        match_results = match_resume_with_jobs(resume_id)
        return jsonify({"results": match_results}), 200

    except Exception as e:
        return jsonify({"error": str(e)}), 500

# 🔹 Run App
if __name__ == "__main__":
    app.run(debug=True)
